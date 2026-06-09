"""Build an integrated GPB manuscript working DOCX from text drafts."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def set_normal_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)


def add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(16)


def add_paragraphs_from_markdown(
    document: Document,
    path: Path,
    *,
    include_heading: bool = True,
    stop_at_heading: str | None = None,
) -> None:
    text = path.read_text(encoding="utf-8", errors="replace")
    in_code = False
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if line.strip() == "```":
            in_code = not in_code
            continue
        if in_code:
            para = document.add_paragraph()
            run = para.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue
        if not line.strip():
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            heading_text = heading.group(2).strip()
            if stop_at_heading and heading_text.lower() == stop_at_heading.lower():
                break
            raw_level = len(heading.group(1))
            if include_heading or raw_level > 1:
                level = min(max(raw_level if include_heading else raw_level - 1, 1), 4)
                document.add_heading(heading_text, level=level)
            continue
        bullet = re.match(r"^-\s+(.*)$", line)
        if bullet:
            document.add_paragraph(bullet.group(1).strip(), style="List Bullet")
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            document.add_paragraph(numbered.group(1).strip(), style="List Number")
            continue
        if line.startswith(">"):
            document.add_paragraph(line.lstrip("> ").strip(), style="Intense Quote")
            continue
        document.add_paragraph(line)


def extract_section(markdown_path: Path, heading: str, next_headings: set[str] | None = None) -> list[str]:
    lines = markdown_path.read_text(encoding="utf-8", errors="replace").splitlines()
    out: list[str] = []
    capture = False
    next_headings = {h.lower() for h in (next_headings or set())}
    for line in lines:
        match = re.match(r"^##\s+(.*)$", line)
        if match:
            text = match.group(1).strip()
            if text.lower() == heading.lower():
                capture = True
                continue
            if capture and (not next_headings or text.lower() in next_headings):
                break
        if capture:
            out.append(line)
    return out


def add_lines(document: Document, lines: list[str]) -> None:
    tmp = "\n".join(lines)
    scratch = Path("__scratch_section__.md")
    scratch.write_text(tmp, encoding="utf-8")
    try:
        add_paragraphs_from_markdown(document, scratch, include_heading=False)
    finally:
        scratch.unlink(missing_ok=True)


def build(args: argparse.Namespace) -> None:
    document = Document()
    set_normal_style(document)

    add_title(document, "anisoNET: Barrier-Aware Tissue-Constrained Spatial Field Inference For Spatial Transcriptomics")
    document.add_paragraph("Working integrated draft generated for GPB revision. Final journal formatting, references, author information, and figure placement remain pending.")

    rewrite = args.text_dir / "manuscript_rewrite_working_draft.md"
    methods = args.text_dir / "methods_integrated_gpb_revision.md"
    legends = args.text_dir / "figure_legends_integrated_gpb_revision.md"

    for heading in ["Abstract", "Introduction", "Results", "Discussion"]:
        document.add_heading(heading, level=1)
        next_heads = {"Introduction", "Results", "Discussion", "Methods Placeholders To Integrate", "Figure Legend Source"}
        add_lines(document, extract_section(rewrite, heading, next_heads))

    document.add_heading("Materials And Methods", level=1)
    add_paragraphs_from_markdown(document, methods, include_heading=False)

    document.add_heading("Figure Legends", level=1)
    add_paragraphs_from_markdown(document, legends, include_heading=False)

    document.add_heading("References", level=1)
    document.add_paragraph("References should be restored from the prior manuscript and updated after final text integration.")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    print(f"Wrote {args.output}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--text-dir", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    build(parser.parse_args())


if __name__ == "__main__":
    main()
