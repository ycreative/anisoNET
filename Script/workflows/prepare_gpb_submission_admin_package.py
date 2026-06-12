"""Prepare GPB submission-administration files from current manuscript assets.

This script does not rerun analyses. It generates conservative submission
documents, locks a supplementary-figure candidate set from existing assets, and
updates reviewer-facing reproducibility package notes for the active v8 draft.
"""

from __future__ import annotations

import csv
import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[2]
TODAY = date(2026, 6, 12).isoformat()
SUBMISSION_DIR = ROOT / "codexAnalysis" / "submission_package"
SUPP_WORKING_DIR = ROOT / "codexAnalysis" / "manuscript_figures" / "supplementary_working"
SUPP_FINAL_DIR = ROOT / "codexAnalysis" / "manuscript_figures" / "supplementary_final_candidate"
REPRO_DIR = ROOT / "codexAnalysis" / "reproducibility_package"
DRAFT_DIR = ROOT / "draft" / "revised"


SUPP_FIGURES = [
    (
        "S1",
        "primary_brain_preflight_qc",
        "Primary brain preprocessing and field-construction QC",
        "Across the eight GSE193107 brain-aging Visium sections, Apoe/CNS-myelin preflight outputs summarize tissue support, measured source, barrier prior, structural resistance, scalar diffusion, and analysis-domain construction. This figure documents that the primary brain cohort was processed by the same field-construction workflow before PINN fitting.",
        "FigS01_primary_brain_preflight_qc_working",
    ),
    (
        "S2",
        "primary_brain_full_fields",
        "Complete Apoe and Gfap primary brain field montages",
        "Full eight-section Apoe and Gfap source-to-field montages support the representative examples and QC summaries in Figure 2. These panels are fitted-field visualizations and should not be described as held-out prediction maps.",
        "FigS02_primary_brain_full_fields_working",
    ),
    (
        "S3",
        "heldout_benchmark_details",
        "Generic held-out interpolation benchmark details",
        "Held-out test summaries and the contextual benchmark panel document the generic interpolation claim boundary. These results are kept separate from fitted-source QC and from the controlled synthetic barrier-leakage mechanism tests.",
        "FigS03_heldout_benchmark_details_working",
    ),
    (
        "S4",
        "synthetic_replicates",
        "Synthetic barrier benchmark replicate maps",
        "Controlled synthetic replicate maps include balanced and train 20%/test 80% stress settings. They provide mechanistic evidence for barrier-aware attenuation under known source, barrier, and ground-truth conditions.",
        "FigS04_synthetic_replicates_working",
    ),
    (
        "S5",
        "histology_prior_sensitivity",
        "Histology-prior sensitivity",
        "Preflight- and PINN-level comparisons evaluate brightness-derived and hematoxylin-enriched structural priors. The figure supports the robustness discussion without claiming that a single histology proxy is universally optimal.",
        "FigS05_histology_prior_sensitivity_working",
    ),
    (
        "S6",
        "profile_and_loss_sensitivity",
        "Loss-profile and low-PDE sensitivity",
        "Multi-section low-PDE and representative loss-weight sensitivity summaries show the source-fit versus smoothness trade-off. The low-PDE profile is a source-fit-optimized sensitivity setting, not a replacement for the conservative default profile.",
        "FigS06_profile_and_loss_sensitivity_working",
    ),
    (
        "S7",
        "parameter_stability",
        "Source clipping, barrier weighting, and seed stability",
        "Source-clipping, alpha-sensitivity, and random-seed summaries document parameter-level stability for the primary brain task. These panels support Figure 4E-F and the robustness-audit narrative.",
        "FigS07_parameter_stability_working",
    ),
    (
        "S8",
        "runtime_memory_convergence",
        "Runtime, memory, and convergence diagnostics",
        "Convergence traces and resource-profile summaries document computational behavior. Runtime and memory are supplementary provenance and are not used as the main Figure 4G result.",
        "FigS08_runtime_memory_convergence_working",
    ),
    (
        "S9",
        "kidney_evidence_boundary",
        "Kidney evidence-boundary analysis",
        "Kidney barrier-split, grid-geodesic, prior-hybrid, and Umod profile-probe outputs summarize non-brain evidence development. The figure supports the claim that barrier-aware priors are task-specific rather than automatically beneficial.",
        "FigS09_kidney_evidence_boundary_working",
    ),
    (
        "S10",
        "kidney_marker_and_targeted_extension",
        "Kidney marker screen and targeted extension",
        "Kidney preflight examples, targeted-field contact sheets, and the marker-extension candidate panel support Figure 5G and Supplementary Table S7. These results document the systematic screen behind the selected kidney marker tasks.",
        "FigS10_kidney_marker_and_targeted_extension_working",
    ),
    (
        "S11",
        "sagittal_brain_negative_control",
        "Sagittal brain negative-control analysis",
        "Sagittal brain preflight and targeted-field outputs show workflow portability in a healthy brain setting where resistance-aware distances do not automatically improve Apoe/Gfap high-barrier prediction. This is retained as a negative-control claim-boundary figure.",
        "FigS11_sagittal_brain_negative_control_working",
    ),
    (
        "S12",
        "liver_apap_task_design",
        "Liver/APAP annotation-aware task design",
        "Liver/APAP preflight, central-marker, annotation-boundary, and targeted-field outputs illustrate why barrier-aware inference depends on biologically aligned source and barrier definitions.",
        "FigS12_liver_apap_task_design_working",
    ),
    (
        "S13",
        "targeted_extension_contact_sheets",
        "Targeted multi-gene extension contact sheets",
        "Full-profile targeted field contact sheets across primary brain, sagittal brain, kidney, and liver/APAP datasets provide visual provenance for the expanded targeted-extension evidence summarized in Figures 3-5.",
        "FigS13_targeted_extension_contact_sheets_working",
    ),
    (
        "S14",
        "marker_module_and_spot_diagnostics",
        "Marker-module and spot-level diagnostics",
        "Leave-one-marker-out summaries and barrier-edge discordant spot examples provide representative checks on marker-module robustness and spot-level behavior near barriers.",
        "FigS14_marker_module_and_spot_diagnostics_working",
    ),
]


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.strip() + "\n", encoding="utf-8")


def markdown_to_docx(markdown: str, out_path: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for raw_line in markdown.strip().splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("> "):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[2:].strip())
            run.italic = True
            paragraph.paragraph_format.left_indent = Inches(0.25)
        else:
            doc.add_paragraph(line)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def cover_letter_text() -> str:
    return f"""
# Cover Letter For GPB Submission

[Date]

Dear Editors,

We are pleased to submit our manuscript, "anisoNET: Barrier-Aware Tissue-Constrained Spatial Field Inference For Spatial Transcriptomics", for consideration as a Research Article in *Genomics, Proteomics & Bioinformatics*.

Spatial transcriptomics preserves tissue architecture, yet many downstream spatial analyses still rely on geometric neighborhoods or smoothing assumptions that do not explicitly encode anatomical barriers. In this manuscript, we present anisoNET as a conservative barrier-aware, tissue-constrained scalar field inference workflow. The method combines Visium spot-level expression, H&E-derived tissue structure, marker-defined transcriptomic barriers, and a physics-informed neural network to infer target-gene fields under source-data, boundary, background, smoothness, and scalar reaction-diffusion residual losses.

The revised evidence package focuses on a primary mouse brain aging Visium application, controlled synthetic barrier simulations, robustness and parameter-sensitivity analyses, and cross-tissue claim-boundary evaluations in kidney, sagittal brain, and liver/APAP datasets. The manuscript uses conservative scalar-model wording and explicit validation boundaries, avoiding broad method-ranking claims beyond the tested tasks.

We believe the manuscript will be of interest to GPB readers because it addresses a practical challenge in spatial omics analysis: how to incorporate tissue barriers and histological context into spatial field interpretation while preserving explicit validation boundaries and reproducibility provenance.

This manuscript has not been published and is not under consideration elsewhere. All authors have approved the submitted version. Conflicts of interest, funding, author contributions, data availability, and code availability statements are included in the manuscript and should be finalized after the author metadata are locked.

Corresponding author:
[AUTHOR TO COMPLETE: name, affiliation, mailing address, email, ORCID if required]

Sincerely,

[AUTHOR TO COMPLETE: corresponding author name]

Preparation note, generated {TODAY}: replace all bracketed author/contact placeholders before submission. Do not use the old cover letter in `draft/Cover Letter.docx`.
"""


def response_letter_text() -> str:
    return f"""
# Response To Reviewers

Manuscript title: anisoNET: Barrier-Aware Tissue-Constrained Spatial Field Inference For Spatial Transcriptomics

Generated template date: {TODAY}

Important submission note: exact editor and reviewer comments must be inserted into the marked slots before resubmission. The response language below is prepared around the current v8 manuscript and Figure 1-5 package, but it is not a complete response letter until the original comments are pasted verbatim.

## Cover Response To The Editor

Dear Editor,

We thank you and the reviewers for the careful evaluation of our manuscript. We have substantially revised the manuscript to clarify the scope, mathematical formulation, validation evidence, figure presentation, supplementary materials, and reproducibility package.

The revised manuscript now presents anisoNET as a barrier-aware, tissue-constrained scalar spatial field inference workflow for spatial transcriptomics. We no longer frame the method as a universal interpolation method, a tensor-valued diffusion model, or a full divergence-form diffusion solver. The revised evidence package separates primary fitted-field reconstruction, generic held-out interpolation claim boundaries, controlled synthetic barrier validation, robustness analyses, and cross-tissue task-boundary evaluations.

Major changes include:

- Rebuilt Figure 1 as a technical method-definition figure describing data inputs, task-specific priors, scalar PINN architecture, loss terms, barrier mechanism, representative output, and validation roadmap.
- Rebuilt Figure 2 around the primary GSE193107 mouse brain aging application with Apoe/Gfap field reconstruction, eight-gene/eight-section QC, tissue-support diagnostics, and barrier-context summaries.
- Reframed Figure 3 to separate generic held-out interpolation from controlled synthetic barrier-leakage validation.
- Added Figure 4 robustness and sensitivity analyses covering histology priors, loss-weight profiles, source clipping, seed stability, and targeted primary-brain extension QC; runtime is retained only as supplementary provenance.
- Added Figure 5 cross-tissue portability and claim-boundary analyses in kidney, sagittal brain, and liver/APAP datasets.
- Finalized a Supplementary Figure S1-S14 candidate set and a Supplementary Table S1-S8 workbook.
- Expanded the GPB-style numbered reference list to 50 cited references.
- Updated the reproducibility package notes for a GitHub + Zenodo submission strategy with an optional reviewer cache for selected intermediate files.

## Response To Reviewer 1

### Comment 1. Mathematical description and model scope

Reviewer comment:

> [[AUTHOR: INSERT EXACT REVIEWER COMMENT 1 HERE BEFORE SUBMISSION]]

Response:

We agree that the mathematical scope needed to be stated more precisely. We revised the manuscript to describe the current two-dimensional implementation as a scalar barrier-aware field model. The implemented residual is written as D(x,y)(C_xx + C_yy) - kC + S(x,y) = 0 and is optimized together with source-data, tissue-domain PDE residual, boundary/background, and smoothness losses. We no longer describe the current implementation as a tensor-valued diffusion model or a full divergence-form div(D grad C) solver. Tensor-valued and divergence-form extensions are now discussed only as future work.

Changes made:

- Revised the Abstract, Introduction, Results, Methods, Discussion, and limitations language.
- Rebuilt Figure 1C around the scalar PINN architecture and scalar residual.
- Added explicit scalar-model caveats to the Figure 1 legend and Methods.
- Ran claim-boundary QC on the integrated v8 manuscript.

### Comment 2. Performance claims and benchmark interpretation

Reviewer comment:

> [[AUTHOR: INSERT EXACT REVIEWER COMMENT 2 HERE BEFORE SUBMISSION]]

Response:

We agree that the previous version implied broader method superiority than the evidence supported. We revised the manuscript to separate generic held-out expression interpolation from the intended barrier-aware field-behavior question. Generic held-out benchmarks are now presented as claim-boundary evidence. Controlled synthetic simulations, where source, barrier, and ground truth are known, are used as the more direct test of barrier-leakage attenuation.

Changes made:

- Rebuilt Figure 3 around held-out claim boundaries and synthetic barrier validation.
- Added leakage and barrier-attenuation summaries across train/test splits and seeds.
- Revised Results and Discussion language to avoid broad superiority wording.
- Added Supplementary Figure S3 and S4 to show held-out details and synthetic replicate maps.

## Response To Reviewer 2

### Comment 3. Robustness and parameter sensitivity

Reviewer comment:

> [[AUTHOR: INSERT EXACT REVIEWER COMMENT 3 HERE BEFORE SUBMISSION]]

Response:

We added a new robustness and sensitivity figure and supporting supplementary material. The revised analyses compare brightness-derived and hematoxylin-enriched histology priors, source clipping thresholds, conservative default and source-fit-optimized low-PDE profiles, random seed stability, and computational cost.

The low-PDE profile improved fitted-source agreement across tested GSE193107 Apoe/Gfap section-target pairs, but it also changed the smoothness trade-off. We therefore report it as a sensitivity profile rather than replacing the conservative default setting.

Changes made:

- Added Figure 4 robustness and sensitivity panels A-G.
- Ensured Figure 4G reports primary brain targeted-extension fitted-source and roughness QC, not runtime.
- Added Supplementary Figures S5-S8 and Supplementary Tables S4 and S8.
- Added Methods details for training profiles, loss weights, and resource profiling.

### Comment 4. Cross-tissue portability and claim boundaries

Reviewer comment:

> [[AUTHOR: INSERT EXACT REVIEWER COMMENT 4 HERE BEFORE SUBMISSION]]

Response:

We revised the cross-tissue interpretation and added analyses in kidney, sagittal brain, and liver/APAP spatial transcriptomics datasets. These datasets are now used to evaluate workflow portability and task boundaries, not to claim broad cross-tissue superiority.

Kidney provides benchmark-development evidence. Healthy sagittal brain serves as a negative-control setting in which resistance-aware distances do not automatically improve Apoe/Gfap high-barrier prediction. Liver/APAP demonstrates that annotation-aware task design is required because barrier-aware modeling depends on biologically aligned source and barrier definitions.

Changes made:

- Added Figure 5 cross-tissue portability and claim-boundary panels.
- Added a kidney systematic marker screen and targeted kidney marker-extension summary.
- Added Supplementary Figures S9-S12 and Supplementary Tables S6-S7.
- Revised the Discussion to emphasize task-specific claim boundaries.

## Response To Reviewer 3

### Comment 5. Figure density, supplementary evidence, and reproducibility

Reviewer comment:

> [[AUTHOR: INSERT EXACT REVIEWER COMMENT 5 HERE BEFORE SUBMISSION]]

Response:

We rebuilt the figure package to make the spatial-transcriptomics evidence more explicit and data-rich. The main figures now prioritize spatial maps, tissue-context panels, compact QC summaries, and evidence matrices. We also organized a final-candidate Supplementary Figure S1-S14 package so that preprocessing, full primary fields, held-out details, synthetic replicate maps, robustness analyses, cross-tissue diagnostics, targeted-extension contact sheets, and spot-level diagnostics are available outside the main figures.

Changes made:

- Finalized Figure 1-5 assembled PDFs and active panel manifests.
- Finalized a Supplementary Figure S1-S14 candidate set with source mapping and legends.
- Updated the supplementary table workbook and supplementary material map.
- Updated the reproducibility package notes to use the v8 manuscript and a GitHub + Zenodo release strategy with optional reviewer cache.

## Remaining Author-Supplied Items Before Submission

- Insert the exact editor and reviewer comments verbatim into the marked slots above.
- Replace author/contact/funding/declaration placeholders in the main manuscript and cover letter.
- Insert final repository URL, Zenodo DOI, or reviewer-cache link when available.
- Run final claim-boundary QC after any manual Word edits.
"""


def supplementary_legends_text() -> str:
    lines = [
        "# Final-Candidate Supplementary Figure Legends",
        "",
        f"Date: {TODAY}",
        "",
        "Status: include Supplementary Figures S1-S14 in the GPB submission package unless the journal portal imposes a file-size limit. These figures are copied from existing generated assets; no model fitting or upstream benchmark was rerun.",
        "",
    ]
    for code, _slug, title, legend, _stem in SUPP_FIGURES:
        lines.extend(
            [
                f"## Supplementary Figure {code}. {title}",
                "",
                legend,
                "",
            ]
        )
    return "\n".join(lines)


def supplementary_set_text(final_entries: list[tuple[str, str, Path, Path]], combined_pdf: Path) -> str:
    lines = [
        "# Final-Candidate Supplementary Figure Set",
        "",
        f"Date: {TODAY}",
        "",
        "Decision: include Supplementary Figures S1-S14 as the current submission candidate set.",
        "",
        "Rationale:",
        "",
        "- S1-S2 support the primary GSE193107 brain-aging application and show the data volume behind Figure 2.",
        "- S3-S4 separate generic held-out interpolation from controlled synthetic barrier validation.",
        "- S5-S8 provide robustness, parameter, convergence, runtime, and memory provenance without using runtime as a main Figure 4G result.",
        "- S9-S12 document kidney, sagittal-brain, and liver/APAP cross-tissue claim boundaries.",
        "- S13-S14 provide full targeted-extension and marker/spot-level diagnostic provenance.",
        "",
        "Files:",
        "",
        f"- Combined PDF: `{combined_pdf.relative_to(ROOT).as_posix()}`",
        "",
    ]
    for code, title, png_path, pdf_path in final_entries:
        rel_png = png_path.relative_to(ROOT).as_posix()
        rel_pdf = pdf_path.relative_to(ROOT).as_posix()
        lines.append(f"- Supplementary Figure {code}: {title}")
        lines.append(f"  - PNG: `{rel_png}`")
        lines.append(f"  - PDF: `{rel_pdf}`")
    lines.extend(
        [
            "",
            "Notes:",
            "",
            "- The `supplementary_working` folder is retained as provenance; use this `supplementary_final_candidate` folder for submission-package assembly.",
            "- The overview file from the working folder is not part of the final candidate set.",
            "- Panel letters are not regenerated inside main-figure panel assets; final figure labels remain assembly-layer labels.",
        ]
    )
    return "\n".join(lines)


def release_decision_text() -> str:
    return f"""
# Submission Release Decision

Date: {TODAY}

Recommended GPB submission strategy:

- Primary code route: public GitHub repository URL.
- Archival route: Zenodo DOI linked to the GitHub release and selected derived materials.
- Reviewer audit route: compact reviewer cache containing final Figure 1-5 panel assets, Supplementary Figure S1-S14 PDFs/PNGs, supplementary table workbook, metric summary tables, and selected representative preflight/PINN arrays.

Use the reviewer cache for bulky intermediate outputs rather than placing full `preflight`, `pinn`, `processed_visium`, or sweep folders in GitHub.

Submission placeholders to replace:

- GitHub repository: `[AUTHOR TO COMPLETE: GitHub repository URL]`
- Zenodo DOI: `[AUTHOR TO COMPLETE: Zenodo DOI or private-review DOI]`
- Reviewer cache: `[AUTHOR TO COMPLETE: reviewer-cache archive path or portal-upload filename]`

Current decision:

- Target a GitHub + Zenodo pair for submission if the repository can be cleaned in time.
- Keep a reviewer cache as a fallback and as a fast-audit supplement.
- Do not promise unrestricted redistribution of raw public datasets; provide accessions and standardization scripts instead.
"""


def reproducibility_readme_text() -> str:
    return f"""
# anisoNET Reproducibility Package Draft

Date: {TODAY}

This folder contains draft public-release materials for the anisoNET GPB submission package. It is intended to be copied or adapted into the root of a clean GitHub repository after final path cleanup and repository metadata are locked.

anisoNET is framed as a barrier-aware, tissue-constrained scalar spatial field inference workflow for spatial transcriptomics. The released implementation should be described as a scalar Laplacian PINN approximation, not as tensor-valued diffusion or full divergence-form diffusion.

## Contents

- `environment.yml`: conda environment specification for the analysis and figure workflows.
- `requirements.txt`: pip-style dependency list for non-conda installs.
- `reproduce_figures.md`: commands and input expectations for regenerating the current Figure 1-5 panel assets and supplementary figure previews from existing outputs.
- `repository_manifest.md`: what belongs in GitHub, Zenodo, and the reviewer cache.
- `data_and_code_availability.md`: manuscript-ready data and code availability draft text with URL/DOI placeholders.
- `zenodo_release_checklist.md`: release checklist for code, derived tables, panel assets, and reviewer cache.
- `submission_release_decision.md`: current decision on GitHub, Zenodo, and reviewer-cache submission routes.

## Minimal Reviewer Workflow

From the repository root, create the environment:

```bash
conda env create -f environment.yml
conda activate anisonet-gpb
```

For the current local workspace, the validated Python interpreter is:

```text
K:\\software\\miniconda\\envs\\scvi_env\\python.exe
```

After public release, scripts should be run from a clean repository root with user-provided `ANISONET_PROJECT_ROOT` or explicit command-line paths. The current Figure 1-5 panel scripts support `ANISONET_PROJECT_ROOT` and `ANISONET_ANALYSIS_ROOT`; some upstream batch scripts still contain local development defaults and should be cleaned before a public release.

## Current Best Submission Files

- Manuscript base: `draft/revised/anisoNET_GPB_integrated_manuscript_v8.docx`
- Cover letter: `draft/revised/anisoNET_GPB_cover_letter_conservative_v1.docx`
- Response template: `draft/revised/anisoNET_GPB_response_to_reviewers_v2_template.docx`
- Supplementary figure legends: `draft/revised/anisoNET_GPB_supplementary_figure_legends_final_candidate.docx`
- Supplementary figures: `codexAnalysis/manuscript_figures/supplementary_final_candidate`
- Supplementary tables: `draft/revised/anisoNET_GPB_supplementary_tables_submission_candidate.xlsx`
- Figure provenance: `codexAnalysis/main_figure_provenance.md`
- Script index: `codexAnalysis/reproducibility_script_index.md`
- References: `codexAnalysis/references/anisoNET_GPB_references_v8.ris`
"""


def data_code_text() -> str:
    return f"""
# Draft Data And Code Availability Text

Date: {TODAY}

## Recommended Submission Text

The source code for anisoNET, including Visium preprocessing utilities, scalar diffusion/resistance field construction, PyTorch-based scalar PINN inference, postprocessing, benchmarking scripts, metric summaries, and figure-panel generation workflows, will be made available at `[AUTHOR TO COMPLETE: GitHub repository URL]` and archived at `[AUTHOR TO COMPLETE: Zenodo DOI]`. A compact reviewer cache containing selected precomputed arrays, metric tables, final panel assets, Supplementary Figure S1-S14 files, and the supplementary table workbook will be provided as `[AUTHOR TO COMPLETE: reviewer-cache archive or portal-upload filename]` if required for peer-review auditing.

The current implementation should be cited as a scalar, tissue-constrained, barrier-aware spatial field inference workflow. It does not implement tensor-valued diffusion or the full divergence-form operator `div(D grad C)`.

## Data Availability

The primary mouse brain aging spatial transcriptomics analysis uses public Visium data from GSE193107. Cross-tissue validation and claim-boundary analyses use public mouse kidney, mouse sagittal brain, and mouse liver/APAP spatial transcriptomics datasets from 10x Genomics public resources and GEO, including GSE280515 where applicable. Accession identifiers, sample identifiers, target genes, barrier markers, and analysis roles are listed in the supplementary dataset inventory table.

Raw public sequencing matrices and histology images are not redistributed in the GitHub repository. Scripts are provided to standardize downloaded public Visium files into the layout expected by the workflow. Lightweight derived metric tables, figure provenance files, current Supplementary Figure S1-S14 assets, and supplementary tables are intended for GitHub/Zenodo release. Larger derived arrays and panel-generation caches should be archived on Zenodo or provided to reviewers as a compact cache when needed for exact figure regeneration.

## Current Decision

Use GitHub + Zenodo as the primary public availability route, and keep a compact reviewer cache as the fast-audit route for large or selected precomputed intermediates.
"""


def zenodo_checklist_text() -> str:
    return f"""
# Zenodo Release Checklist

Date: {TODAY}

## Before Creating The Release

- [x] Active submission manuscript recorded as v8.
- [x] Conservative cover letter prepared.
- [x] Response-to-reviewers template updated; exact reviewer comments still require author insertion.
- [x] Final-candidate Supplementary Figure S1-S14 set prepared from existing assets.
- [x] Supplementary table workbook exists.
- [x] Figure 1-5 final PDFs exist in figure folders.
- [x] Figure 1-5 panel-generation scripts support repository-root inference and environment-variable overrides.
- [x] Sanitized dataset example config exists.
- [ ] Replace author/contact/funding/declaration placeholders.
- [ ] Create or confirm public GitHub repository URL.
- [ ] Create Zenodo draft record and DOI.
- [ ] Add license file.
- [ ] Add `CITATION.cff`.
- [ ] Add package-level public `README.md`.
- [ ] Add synthetic smoke test.
- [ ] Run `py_compile` on selected public `Script/anisonet` and cleaned `Script/workflows` files.
- [ ] Confirm final `CURRENT_PANEL_SET.md` files match submitted figures.
- [ ] Re-export supplementary table workbook after any final edits.
- [ ] Run final manuscript/response/cover claim-boundary QC.

## Archive Components

- [ ] Code source archive.
- [ ] Derived metric summary tables.
- [x] Supplementary table workbook candidate.
- [x] Main figure panel assets and final assembled PDFs exist locally.
- [x] Supplementary Figure S1-S14 candidate assets exist locally.
- [ ] Reviewer cache with selected arrays and JSON metadata.
- [ ] README describing fast audit versus full rerun routes.

## Metadata

- [ ] Title uses `anisoNET`.
- [ ] Authors match manuscript.
- [ ] Related identifiers include GitHub release and GPB manuscript if available.
- [ ] Keywords include spatial transcriptomics, physics-informed neural network, anatomical barriers, Visium, reproducibility.
- [ ] Version tag matches GitHub release.
- [ ] License is consistent across GitHub and Zenodo.
"""


def submission_status_text() -> str:
    return f"""
# Submission Administrative Package Status

Date: {TODAY}

Completed in this pass:

- Prepared a conservative GPB cover letter and Word export.
- Rebuilt the response-to-reviewers template around the current v8 manuscript and Figure 1-5 package.
- Replaced the stale reviewer-comment placeholder phrase with explicit author-action slots for exact reviewer comments.
- Locked a final-candidate Supplementary Figure S1-S14 set from existing generated assets.
- Prepared Supplementary Figure S1-S14 legends and source map.
- Updated reproducibility package notes from the old v4 manuscript reference to the current v8 manuscript.
- Chose GitHub + Zenodo as the primary submission availability route, with a compact reviewer cache as the fallback/fast-audit route.
- Cover letter claim-boundary QC has 0 hits after conservative wording cleanup.
- Response-template claim-boundary QC has expected limitation/removal-context hits because it explicitly states which older claims were narrowed.

Still author-dependent:

- Author list, affiliations, corresponding author email, ORCID IDs, funding, acknowledgements, author contributions, and conflict-of-interest statements.
- Exact editor/reviewer comments must be pasted verbatim into the response letter.
- Final GitHub URL, Zenodo DOI, and reviewer-cache archive name or link.

Do not use:

- `draft/Cover Letter.docx` for submission; it is an old overclaiming cover letter.
- `codexAnalysis/manuscript_figures/supplementary_working` as the final supplementary folder; it remains provenance only.
"""


def copy_supplementary_figures() -> list[tuple[str, str, Path, Path]]:
    SUPP_FINAL_DIR.mkdir(parents=True, exist_ok=True)
    entries: list[tuple[str, str, Path, Path]] = []
    source_rows: list[dict[str, str]] = []
    source_map_path = SUPP_WORKING_DIR / "supplementary_working_source_map.csv"
    if source_map_path.exists():
        with source_map_path.open(newline="", encoding="utf-8") as handle:
            source_rows = list(csv.DictReader(handle))

    with (SUPP_FINAL_DIR / "supplementary_figure_source_map_final.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["supplementary_figure", "title", "final_png", "final_pdf", "working_stem", "source_label", "source_path"])
        for code, slug, title, _legend, stem in SUPP_FIGURES:
            src_png = SUPP_WORKING_DIR / f"{stem}.png"
            src_pdf = SUPP_WORKING_DIR / f"{stem}.pdf"
            final_png = SUPP_FINAL_DIR / f"Supplementary_Figure_{code}_{slug}.png"
            final_pdf = SUPP_FINAL_DIR / f"Supplementary_Figure_{code}_{slug}.pdf"
            if not src_png.exists() or not src_pdf.exists():
                raise FileNotFoundError(f"Missing supplementary source files for {code}: {src_png} / {src_pdf}")
            shutil.copy2(src_png, final_png)
            shutil.copy2(src_pdf, final_pdf)
            entries.append((code, title, final_png, final_pdf))
            matched = [row for row in source_rows if row.get("figure") == f"Fig{code}"]
            if not matched:
                writer.writerow([code, title, final_png.relative_to(ROOT).as_posix(), final_pdf.relative_to(ROOT).as_posix(), stem, "", ""])
            for row in matched:
                writer.writerow(
                    [
                        code,
                        title,
                        final_png.relative_to(ROOT).as_posix(),
                        final_pdf.relative_to(ROOT).as_posix(),
                        stem,
                        row.get("label", ""),
                        row.get("source_path", ""),
                    ]
                )
    return entries


def build_combined_supplementary_pdf(final_entries: list[tuple[str, str, Path, Path]]) -> Path:
    combined_pdf = SUPP_FINAL_DIR / "Supplementary_Figures_S1-S14_final_candidate.pdf"
    pages: list[Image.Image] = []
    for _code, _title, png_path, _pdf_path in final_entries:
        pages.append(Image.open(png_path).convert("RGB"))
    if not pages:
        raise RuntimeError("No supplementary figure pages found for combined PDF")
    first, rest = pages[0], pages[1:]
    first.save(combined_pdf, "PDF", resolution=220.0, save_all=True, append_images=rest)
    for page in pages:
        page.close()
    return combined_pdf


def main() -> None:
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    DRAFT_DIR.mkdir(parents=True, exist_ok=True)
    REPRO_DIR.mkdir(parents=True, exist_ok=True)

    cover_md = cover_letter_text()
    response_md = response_letter_text()
    legends_md = supplementary_legends_text()
    release_md = release_decision_text()

    write_text(SUBMISSION_DIR / "cover_letter_gpb_conservative_v1.md", cover_md)
    write_text(SUBMISSION_DIR / "response_to_reviewers_template_v2.md", response_md)
    write_text(ROOT / "codexAnalysis" / "response_to_reviewers_working_draft.md", response_md)
    write_text(SUBMISSION_DIR / "supplementary_figure_legends_final_candidate.md", legends_md)
    write_text(REPRO_DIR / "submission_release_decision.md", release_md)
    write_text(REPRO_DIR / "README.md", reproducibility_readme_text())
    write_text(REPRO_DIR / "data_and_code_availability.md", data_code_text())
    write_text(REPRO_DIR / "zenodo_release_checklist.md", zenodo_checklist_text())
    write_text(SUBMISSION_DIR / "submission_admin_status_20260612.md", submission_status_text())

    markdown_to_docx(cover_md, DRAFT_DIR / "anisoNET_GPB_cover_letter_conservative_v1.docx")
    markdown_to_docx(response_md, DRAFT_DIR / "anisoNET_GPB_response_to_reviewers_v2_template.docx")
    markdown_to_docx(response_md, DRAFT_DIR / "response_to_reviewers_working_draft.docx")
    markdown_to_docx(response_md, DRAFT_DIR / "anisoNET_GPB_response_to_reviewers_working.docx")
    markdown_to_docx(legends_md, DRAFT_DIR / "anisoNET_GPB_supplementary_figure_legends_final_candidate.docx")

    final_entries = copy_supplementary_figures()
    combined_pdf = build_combined_supplementary_pdf(final_entries)
    write_text(SUPP_FINAL_DIR / "CURRENT_SUPPLEMENTARY_FIGURE_SET.md", supplementary_set_text(final_entries, combined_pdf))
    write_text(SUPP_FINAL_DIR / "supplementary_figure_legends_final_candidate.md", legends_md)

    print("Prepared GPB submission administration package.")
    print(f"Submission docs: {SUBMISSION_DIR}")
    print(f"Supplementary figures: {SUPP_FINAL_DIR}")
    print(f"Reproducibility docs: {REPRO_DIR}")


if __name__ == "__main__":
    main()
