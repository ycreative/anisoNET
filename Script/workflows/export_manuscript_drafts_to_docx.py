"""Export manuscript Markdown drafts to Word DOCX working files.

The converter is intentionally simple: it preserves headings, bullets,
numbered lists, paragraphs, and code-like lines well enough for Word review.
Final GPB formatting should still be done manually in Word.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_code_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_markdown_line(document: Document, line: str, in_code: bool) -> bool:
    stripped = line.rstrip()
    if stripped.strip() == "```":
        return not in_code
    if in_code:
        add_code_paragraph(document, stripped)
        return in_code
    if not stripped:
        return in_code

    heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if heading:
        level = min(len(heading.group(1)), 4)
        document.add_heading(heading.group(2).strip(), level=level)
        return in_code

    bullet = re.match(r"^-\s+(.*)$", stripped)
    if bullet:
        document.add_paragraph(bullet.group(1).strip(), style="List Bullet")
        return in_code

    numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
    if numbered:
        document.add_paragraph(numbered.group(1).strip(), style="List Number")
        return in_code

    if stripped.startswith(">"):
        paragraph = document.add_paragraph(stripped.lstrip("> ").strip())
        paragraph.style = document.styles["Intense Quote"]
        return in_code

    document.add_paragraph(stripped)
    return in_code


def export_docx(input_path: Path, output_path: Path) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    text = input_path.read_text(encoding="utf-8", errors="replace")
    in_code = False
    for line in text.splitlines():
        in_code = add_markdown_line(document, line, in_code)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--outdir", required=True, type=Path)
    parser.add_argument("inputs", nargs="+", type=Path)
    args = parser.parse_args()

    for input_path in args.inputs:
        output_path = args.outdir / f"{input_path.stem}.docx"
        export_docx(input_path, output_path)
        print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
